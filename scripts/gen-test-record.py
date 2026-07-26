"""Generate Test Record Document (Word) - test execution log template."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = os.path.join(os.path.dirname(__file__), '..', 'docs', '测试记录文档.docx')

# ── Colors ──────────────────────────────────────────────
PRIMARY = RGBColor(0x1a, 0x5c, 0x4b)
ACCENT = RGBColor(0xc0, 0x5f, 0x2a)
MUTED = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xc0, 0x39, 0x2b)
GREEN = RGBColor(0x2d, 0x7a, 0x4f)
TABLE_HEADER_BG = 'D5E8D4'
TABLE_ALT_BG = 'F5F5F5'

# ── Helpers ─────────────────────────────────────────────
def set_cell_shading(cell, fill_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): fill_color
    })
    tc_pr.append(shd)

def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn('w:tcBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:color'): '999999'
        })
        borders.append(el)
    tc_pr.append(borders)

def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = PRIMARY
        set_cell_shading(cell, TABLE_HEADER_BG)
        set_cell_border(cell)
    # body
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            set_cell_border(cell)
    return table

def add_body(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
    return h

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_result_row(table, row_idx, test_id, module, url, result, note):
    """Fill a test result row with color-coded result."""
    vals = [test_id, module, url, result, note]
    for ci, val in enumerate(vals):
        cell = table.rows[row_idx].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 4 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.5)
                if ci == 3:  # result column
                    if result == '通过':
                        run.font.color.rgb = GREEN
                        run.bold = True
                    elif result == '失败':
                        run.font.color.rgb = RED
                        run.bold = True
                    elif result == '阻塞':
                        run.font.color.rgb = ACCENT
                        run.bold = True
        set_cell_border(cell)

def add_blank_result_rows(doc, headers, data_rows, col_widths):
    """Create a result table with pre-filled test IDs and blank result columns."""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = PRIMARY
        set_cell_shading(cell, TABLE_HEADER_BG)
        set_cell_border(cell)
    # data rows
    for ri, (test_id, module, url) in enumerate(data_rows):
        row_data = [test_id, module, url, '', '']
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 4 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if ri % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            set_cell_border(cell)
    return table

# ═════════════════════════════════════════════════════════════
# Document
# ═════════════════════════════════════════════════════════════
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── Cover ──────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('测试记录文档')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('人工智能训练师五级练习与考试系统')
run.font.size = Pt(16)
run.font.color.rgb = MUTED

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('版本：V1.0')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('日期：________年______月______日')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('测试人员：________________')
run.font.size = Pt(12)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 一、测试基本信息
# ═════════════════════════════════════════════════════════════

add_heading(doc, '一、测试基本信息', level=1)

add_table(doc,
    ['项目', '内容'],
    [
        ['系统名称', '人工智能训练师五级练习与考试系统'],
        ['版本号', 'V1.0'],
        ['测试地址', '________________________________________'],
        ['测试日期', '______年____月____日 至 ______年____月____日'],
        ['测试人员', '________________________________________'],
        ['浏览器', '________________________________________'],
        ['测试范围', '登录认证 + 学员端 + 教师端 + 管理端全功能'],
    ],
    col_widths_cm=[4, 11])

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════
# 二、测试结果汇总表
# ═════════════════════════════════════════════════════════════

add_heading(doc, '二、测试结果汇总表', level=1)

add_body(doc, '说明：在「结果」列填写「通过」「失败」或「阻塞」，在「备注」列记录问题摘要或截图编号。', color=MUTED, size=10)

# ── 2.1 登录认证 ───────────────────────────────────────
add_heading(doc, '2.1  登录认证模块', level=2)

login_tests = [
    ('TC-AUTH-01', '学员正常登录', '/login'),
    ('TC-AUTH-02', '错误密码登录', '/login'),
    ('TC-AUTH-03', '未登录访问受保护页面', '/student/home'),
    ('TC-AUTH-04', '角色权限隔离', '/admin/dashboard'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    login_tests,
    [2.5, 4, 4, 2, 3])

doc.add_paragraph()

# ── 2.2 学员端-理论练习 ───────────────────────────────
add_heading(doc, '2.2  学员端 - 理论练习', level=2)

practice_tests = [
    ('TC-PRACTICE-01', '进入练习并答题', '/student/practice'),
    ('TC-PRACTICE-02', '错题自动收集', '/student/practice → /student/wrong'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    practice_tests,
    [2.5, 4, 4, 2, 3])

doc.add_paragraph()

# ── 2.3 学员端-实操任务 ───────────────────────────────
add_heading(doc, '2.3  学员端 - 实操任务', level=2)

task_tests = [
    ('TC-TASK-01', '图片清洗任务', '/student/task'),
    ('TC-TASK-02', '矩形框标注任务', '/student/task'),
    ('TC-TASK-03', '点标注任务', '/student/task'),
    ('TC-TASK-04', '文本情感标注任务', '/student/task'),
    ('TC-TASK-05', '音频转写任务', '/student/task'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    task_tests,
    [2.5, 4, 4, 2, 3])

doc.add_paragraph()

# ── 2.4 学员端-考试模块 ───────────────────────────────
add_heading(doc, '2.4  学员端 - 考试模块', level=2)

exam_tests = [
    ('TC-EXAM-01', '开始考试', '/student/exams'),
    ('TC-EXAM-02', '考试中自动保存', '/student/exams'),
    ('TC-EXAM-03', '手动交卷', '/student/exams'),
    ('TC-EXAM-04', '超时自动交卷', '/student/exams'),
    ('TC-EXAM-05', '未到考试时间不可开始', '/student/exams'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    exam_tests,
    [2.5, 4, 4, 2, 3])

doc.add_paragraph()

# ── 2.5 学员端-成绩查询 ───────────────────────────────
add_heading(doc, '2.5  学员端 - 成绩查询', level=2)

result_tests = [
    ('TC-RESULT-01', '查看已发布成绩', '/student/results'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    result_tests,
    [2.5, 4, 4, 2, 3])

doc.add_paragraph()

# ── 2.6 教师端 ───────────────────────────────────────
add_heading(doc, '2.6  教师端模块', level=2)

teacher_tests = [
    ('TC-TEACHER-01', '教师仪表盘', '/teacher/dashboard'),
    ('TC-TEACHER-02', '教师查看学员列表', '/teacher/students'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    teacher_tests,
    [2.5, 4, 4, 2, 3])

doc.add_paragraph()

# ── 2.7 管理端 ───────────────────────────────────────
add_heading(doc, '2.7  管理端模块', level=2)

admin_tests = [
    ('TC-ADMIN-01', '管理端工作台', '/admin/dashboard'),
    ('TC-ADMIN-02', '成绩复核-查看详情', '/admin/results'),
    ('TC-ADMIN-03', '成绩复核-确认发布', '/admin/results'),
    ('TC-ADMIN-04', '成绩复核-手动调整分数', '/admin/results'),
    ('TC-ADMIN-05', '成绩复核-调整校验', '/admin/results'),
    ('TC-ADMIN-06', '题库导入', '/admin/import'),
    ('TC-ADMIN-07', '考务安排', '/admin/exam-schedules'),
    ('TC-ADMIN-08', '审计日志查询', '/admin/audit'),
]
add_blank_result_rows(doc,
    ['用例编号', '测试模块', '页面 URL', '结果', '备注'],
    admin_tests,
    [2.5, 4, 4, 2, 3])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 三、缺陷记录表
# ═════════════════════════════════════════════════════════════

add_heading(doc, '三、缺陷记录表', level=1)

add_body(doc, '说明：每发现一个问题，复制下方缺陷模板填写。截图直接粘贴到「问题截图」区域。', color=MUTED, size=10)

# ── 缺陷模板 × 5 ──────────────────────────────────────
for i in range(1, 6):
    add_heading(doc, f'缺陷 #{i}（BUG-2026-{i:03d}）', level=2)

    add_table(doc,
        ['字段', '内容'],
        [
            ['Bug ID', f'BUG-2026-{i:03d}'],
            ['发现日期', '______年____月____日'],
            ['发现人', '________________'],
            ['测试角色/账号', '□ 学员  □ 教师  □ 管理员    账号：________________'],
            ['页面 URL', '________________________________________'],
            ['严重等级', '□ S1 致命   □ S2 严重   □ S3 一般   □ S4 轻微'],
            ['Bug 标题', '________________________________________'],
            ['问题描述', '预期结果：\n\n实际结果：\n'],
            ['复现步骤', '1.\n2.\n3.\n'],
            ['问题截图', '\n\n（在此处粘贴截图）\n\n'],
            ['状态', '□ 新建   □ 已确认   □ 已修复   □ 已关闭'],
        ],
        col_widths_cm=[3.5, 11.5])

    doc.add_paragraph()

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 四、测试统计汇总
# ═════════════════════════════════════════════════════════════

add_heading(doc, '四、测试统计汇总', level=1)

add_heading(doc, '4.1  用例执行统计', level=2)
add_table(doc,
    ['统计项', '数量', '占比'],
    [
        ['测试用例总数', '27', '100%'],
        ['通过', '______', '______%'],
        ['失败', '______', '______%'],
        ['阻塞（环境/数据问题）', '______', '______%'],
        ['未执行', '______', '______%'],
    ],
    col_widths_cm=[6, 4, 4])

add_heading(doc, '4.2  缺陷统计', level=2)
add_table(doc,
    ['严重等级', '数量', '已修复', '未关闭'],
    [
        ['S1 致命', '______', '______', '______'],
        ['S2 严重', '______', '______', '______'],
        ['S3 一般', '______', '______', '______'],
        ['S4 轻微', '______', '______', '______'],
        ['合计', '______', '______', '______'],
    ],
    col_widths_cm=[4, 3, 3, 3])

add_heading(doc, '4.3  测试结论', level=2)
add_table(doc,
    ['结论项', '选择'],
    [
        ['是否达到上线标准', '□ 是，建议上线   □ 否，需修复后回归   □ 有条件上线'],
        ['遗留风险说明', '\n\n\n'],
        ['测试人员签字', '________________'],
        ['日期', '______年____月____日'],
    ],
    col_widths_cm=[5, 10])

# ── save ─────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Test record saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
